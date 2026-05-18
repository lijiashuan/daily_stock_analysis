# -*- coding: utf-8 -*-
"""
报告导出服务 - 支持多种格式导出

支持的格式:
- Markdown (.md)
- Word Document (.docx)
- Rich Text Format (.rtf)
- HTML (.html)
- PDF (.pdf)
"""
import logging
from pathlib import Path
from typing import Optional
from datetime import datetime

logger = logging.getLogger(__name__)


class ReportExportService:
    """报告导出服务,支持多种格式"""

    @staticmethod
    def export_to_docx(content: str, filepath: Optional[Path] = None) -> str:
        """
        将 Markdown 内容导出为 Word 文档

        Args:
            content: Markdown 格式的报告内容
            filepath: 输出文件路径(可选)

        Returns:
            保存的文件路径
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import re

            if filepath is None:
                reports_dir = Path(__file__).parent.parent.parent / 'reports'
                reports_dir.mkdir(parents=True, exist_ok=True)
                date_str = datetime.now().strftime('%Y%m%d')
                filepath = reports_dir / f"report_{date_str}.docx"

            doc = Document()

            # 设置默认样式
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Microsoft YaHei'
            font.size = Pt(11)

            # 解析 Markdown 并转换为 Word
            lines = content.split('\n')
            i = 0

            while i < len(lines):
                line = lines[i].strip()

                # 标题处理
                if line.startswith('# '):
                    heading = doc.add_heading(line[2:], level=1)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in heading.runs:
                        run.font.name = 'Microsoft YaHei'
                        run.font.size = Pt(16)
                        run.font.bold = True

                elif line.startswith('## '):
                    heading = doc.add_heading(line[3:], level=2)
                    for run in heading.runs:
                        run.font.name = 'Microsoft YaHei'
                        run.font.size = Pt(14)
                        run.font.bold = True

                elif line.startswith('### '):
                    heading = doc.add_heading(line[4:], level=3)
                    for run in heading.runs:
                        run.font.name = 'Microsoft YaHei'
                        run.font.size = Pt(12)
                        run.font.bold = True

                # 引用块
                elif line.startswith('> '):
                    p = doc.add_paragraph(line[2:])
                    p_format = p.paragraph_format
                    p_format.left_indent = Pt(20)
                    for run in p.runs:
                        run.font.name = 'Microsoft YaHei'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(100, 100, 100)

                # 分隔线
                elif line == '---':
                    doc.add_paragraph('_' * 50)

                # 表格检测 - 支持完整的表格解析
                elif line.startswith('|'):
                    # 收集表格的所有行
                    table_rows = []
                    while i < len(lines) and lines[i].strip().startswith('|'):
                        row_line = lines[i].strip()
                        # 解析表格行：去除首尾的 | 并按 | 分割
                        cells = [cell.strip() for cell in row_line.split('|') if cell.strip()]
                        # 跳过表格分隔行（全是 --- 的行）
                        if cells and all(c.startswith('---') for c in cells):
                            i += 1
                            continue
                        table_rows.append(cells)
                        i += 1
                    
                    # 如果收集到了表格数据，创建 Word 表格
                    if table_rows:
                        # 检查表格是否完整（每行单元格数一致）
                        max_cols = max(len(row) for row in table_rows)
                        # 补齐单元格
                        for row in table_rows:
                            while len(row) < max_cols:
                                row.append('')
                        
                        # 创建表格（第一行作为表头）
                        if len(table_rows) >= 2:
                            word_table = doc.add_table(rows=len(table_rows), cols=max_cols)
                            word_table.style = 'Table Grid'
                            
                            for row_idx, cells in enumerate(table_rows):
                                for col_idx, cell_text in enumerate(cells):
                                    cell = word_table.cell(row_idx, col_idx)
                                    cell.text = cell_text
                                    # 设置表头样式
                                    if row_idx == 0:
                                        for run in cell.paragraphs[0].runs:
                                            run.font.name = 'Microsoft YaHei'
                                            run.font.size = Pt(10)
                                            run.font.bold = True
                                    else:
                                        for run in cell.paragraphs[0].runs:
                                            run.font.name = 'Microsoft YaHei'
                                            run.font.size = Pt(10)
                            
                            doc.add_paragraph()  # 表格后加空行
                            continue
                        else:
                            # 只有一行的表格，当作普通文本处理
                            doc.add_paragraph(' | '.join(table_rows[0]))
                            continue
                    else:
                        # 如果表格解析失败，继续处理当前行
                        pass

                # 列表项
                elif line.startswith('- ') or line.startswith('* '):
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                    for run in p.runs:
                        run.font.name = 'Microsoft YaHei'
                        run.font.size = Pt(11)

                # 普通段落
                elif line:
                    # 处理粗体和斜体
                    p = doc.add_paragraph()
                    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
                    for part in parts:
                        if not part:
                            continue
                        run = p.add_run(part)
                        run.font.name = 'Microsoft YaHei'
                        run.font.size = Pt(11)

                        if part.startswith('**') and part.endswith('**'):
                            run.bold = True
                            run.text = part[2:-2]
                        elif part.startswith('*') and part.endswith('*'):
                            run.italic = True
                            run.text = part[1:-1]

                i += 1

            doc.save(str(filepath))
            logger.info(f"Word 文档已保存到: {filepath}")
            return str(filepath)

        except ImportError:
            logger.error("缺少 python-docx 库,请运行: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"导出 Word 文档失败: {e}", exc_info=True)
            raise

    @staticmethod
    def export_to_rtf(content: str, filepath: Optional[Path] = None) -> str:
        """
        将 Markdown 内容导出为 RTF 富文本格式

        Args:
            content: Markdown 格式的报告内容
            filepath: 输出文件路径(可选)

        Returns:
            保存的文件路径
        """
        try:
            if filepath is None:
                reports_dir = Path(__file__).parent.parent.parent / 'reports'
                reports_dir.mkdir(parents=True, exist_ok=True)
                date_str = datetime.now().strftime('%Y%m%d')
                filepath = reports_dir / f"report_{date_str}.rtf"

            # RTF 文档头 - 使用 UTF-8 编码声明
            rtf_content = [
                r"{\rtf1\utf8\deff0",
                r"{\fonttbl{\f0\fswiss\fcharset0 Microsoft YaHei;}}",
                r"\viewkind4\uc1\pard",
                r"\f0\fs22",  # 字体大小 11pt
                ""
            ]

            def escape_rtf_text(text: str) -> str:
                """转义 RTF 特殊字符并处理 Unicode"""
                # 转义 RTF 特殊字符
                text = text.replace('\\', '\\\\')
                text = text.replace('{', '\\{')
                text = text.replace('}', '\\}')
                
                # 处理非 ASCII 字符（中文等）
                result = []
                for char in text:
                    if ord(char) > 127:
                        # Unicode 字符使用 \u 转义
                        result.append(f"\\u{ord(char)}?")
                    else:
                        result.append(char)
                return ''.join(result)

            lines = content.split('\n')

            for line in lines:
                line = line.strip()

                if not line:
                    rtf_content.append("\\par")
                    continue

                # 标题
                if line.startswith('# '):
                    text = escape_rtf_text(line[2:])
                    rtf_content.append(f"\\pard\\sb240\\sa120\\b\\fs32 {text}\\b0\\par")

                elif line.startswith('## '):
                    text = escape_rtf_text(line[3:])
                    rtf_content.append(f"\\pard\\sb180\\sa100\\b\\fs28 {text}\\b0\\par")

                elif line.startswith('### '):
                    text = escape_rtf_text(line[4:])
                    rtf_content.append(f"\\pard\\sb120\\sa80\\b\\fs24 {text}\\b0\\par")

                # 引用
                elif line.startswith('> '):
                    text = escape_rtf_text(line[2:])
                    rtf_content.append(f"\\pard\\li360\\ri360\\cf1\\fs20 {text}\\cf0\\par")

                # 分隔线
                elif line == '---':
                    rtf_content.append("\\pard\\brdrb\\brdrs\\brdrw10\\brsp40 \\par")

                # 列表
                elif line.startswith('- ') or line.startswith('* '):
                    text = escape_rtf_text(line[2:])
                    rtf_content.append(f"\\pard\\li360\\bullet {text}\\par")

                # 普通段落
                else:
                    # 简单的粗体处理
                    text = escape_rtf_text(line)
                    text = text.replace('**', '\\b ').replace('**', '\\b0 ')
                    rtf_content.append(f"\\pard\\sa60 {text}\\par")

            rtf_content.append("}")

            # 使用 UTF-8 编码保存
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write('\n'.join(rtf_content))

            logger.info(f"RTF 文档已保存到: {filepath}")
            return str(filepath)

        except Exception as e:
            logger.error(f"导出 RTF 文档失败: {e}", exc_info=True)
            raise

    @staticmethod
    def export_to_html(content: str, filepath: Optional[Path] = None) -> str:
        """
        将 Markdown 内容导出为 HTML

        Args:
            content: Markdown 格式的报告内容
            filepath: 输出文件路径(可选)

        Returns:
            保存的文件路径
        """
        try:
            import markdown2

            if filepath is None:
                reports_dir = Path(__file__).parent.parent.parent / 'reports'
                reports_dir.mkdir(parents=True, exist_ok=True)
                date_str = datetime.now().strftime('%Y%m%d')
                filepath = reports_dir / f"report_{date_str}.html"

            # 转换为 HTML
            html_body = markdown2.markdown(
                content,
                extras=['tables', 'fenced-code-blocks', 'header-ids']
            )

            # 完整的 HTML 文档
            html_doc = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>股票分析报告</title>
    <style>
        body {{
            font-family: "Microsoft YaHei", Arial, sans-serif;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
            color: #333;
        }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #34495e; margin-top: 30px; }}
        h3 {{ color: #7f8c8d; }}
        blockquote {{
            border-left: 4px solid #3498db;
            padding-left: 20px;
            margin-left: 0;
            color: #7f8c8d;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 8px 12px;
            text-align: left;
        }}
        th {{ background-color: #3498db; color: white; }}
        tr:nth-child(even) {{ background-color: #f2f2f2; }}
        hr {{ border: none; border-top: 1px solid #ddd; margin: 30px 0; }}
        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: monospace;
        }}
        pre {{
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
        }}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""

            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(html_doc)

            logger.info(f"HTML 文档已保存到: {filepath}")
            return str(filepath)

        except ImportError:
            logger.error("缺少 markdown2 库,已在 requirements.txt 中")
            raise
        except Exception as e:
            logger.error(f"导出 HTML 文档失败: {e}", exc_info=True)
            raise

    @staticmethod
    def export_to_pdf(content: str, filepath: Optional[Path] = None) -> str:
        """
        将 Markdown 内容导出为 PDF 文档

        Args:
            content: Markdown 格式的报告内容
            filepath: 输出文件路径(可选)

        Returns:
            保存的文件路径
        """
        try:
            import markdown2

            if filepath is None:
                reports_dir = Path(__file__).parent.parent.parent / 'reports'
                reports_dir.mkdir(parents=True, exist_ok=True)
                date_str = datetime.now().strftime('%Y%m%d')
                filepath = reports_dir / f"report_{date_str}.pdf"

            # 将 Markdown 转换为 HTML
            html_body = markdown2.markdown(
                content,
                extras=['tables', 'fenced-code-blocks', 'header-ids']
            )

            # 完整的 HTML 文档
            html_doc = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>股票分析报告</title>
    <style>
        body {{
            font-family: "Microsoft YaHei", "SimSun", Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            font-size: 11pt;
            margin: 2cm;
        }}
        h1 {{ 
            color: #2c3e50; 
            border-bottom: 2px solid #3498db; 
            padding-bottom: 10px;
            font-size: 18pt;
        }}
        h2 {{ 
            color: #34495e; 
            margin-top: 20px;
            font-size: 15pt;
        }}
        h3 {{ 
            color: #7f8c8d;
            font-size: 13pt;
        }}
        blockquote {{
            border-left: 4px solid #3498db;
            padding-left: 20px;
            margin-left: 0;
            color: #7f8c8d;
            background-color: #f8f9fa;
            padding: 10px 20px;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 15px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 8px 12px;
            text-align: left;
        }}
        th {{ 
            background-color: #3498db; 
            color: white;
            font-weight: bold;
        }}
        tr:nth-child(even) {{ 
            background-color: #f2f2f2; 
        }}
        hr {{ 
            border: none; 
            border-top: 1px solid #ddd; 
            margin: 20px 0; 
        }}
        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: "Courier New", monospace;
            font-size: 10pt;
        }}
        pre {{
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
        }}
        ul, ol {{
            margin: 10px 0;
            padding-left: 30px;
        }}
        li {{
            margin: 5px 0;
        }}
        p {{
            margin: 8px 0;
            text-align: justify;
        }}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""

            # 尝试使用 wkhtmltopdf 直接生成 PDF
            try:
                import subprocess
                
                # Windows 上硬编码 wkhtmltopdf 路径
                wkhtmltopdf_path = r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe'
                
                # 检查 wkhtmltopdf 是否存在
                if not Path(wkhtmltopdf_path).exists():
                    # 尝试从 PATH 中查找
                    import shutil
                    wkhtmltopdf_bin = shutil.which('wkhtmltopdf')
                    if wkhtmltopdf_bin:
                        wkhtmltopdf_path = wkhtmltopdf_bin
                    else:
                        raise OSError("wkhtmltopdf not found")
                
                # 使用 subprocess 直接调用 wkhtmltopdf
                cmd = [
                    wkhtmltopdf_path,
                    '--page-size', 'A4',
                    '--margin-top', '20mm',
                    '--margin-right', '20mm',
                    '--margin-bottom', '20mm',
                    '--margin-left', '20mm',
                    '--encoding', 'UTF-8',
                    '--no-outline',
                    '--enable-local-file-access',
                    '-',  # 从 stdin 读取 HTML
                    str(filepath)  # 输出文件
                ]
                
                # 执行命令
                process = subprocess.run(
                    cmd,
                    input=html_doc.encode('utf-8'),
                    capture_output=True,
                    timeout=60
                )
                
                if process.returncode != 0:
                    error_msg = process.stderr.decode('utf-8', errors='ignore')
                    raise RuntimeError(f"wkhtmltopdf failed: {error_msg}")
                
                logger.info(f"PDF 文档已保存到 (wkhtmltopdf): {filepath}")
                return str(filepath)
                
            except FileNotFoundError:
                logger.warning("wkhtmltopdf 未找到，尝试使用 weasyprint...")
                raise

        except (ImportError, OSError):
            # 如果 imgkit 失败，尝试 weasyprint
            try:
                from weasyprint import HTML, CSS
                
                # 重新生成 HTML（因为上面可能在 try 块中）
                import markdown2
                html_body = markdown2.markdown(
                    content,
                    extras=['tables', 'fenced-code-blocks', 'header-ids']
                )
                
                html_doc = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>股票分析报告</title>
    <style>
        @page {{
            size: A4;
            margin: 2cm;
        }}
        body {{
            font-family: "Microsoft YaHei", "SimSun", Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            font-size: 11pt;
        }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; font-size: 18pt; }}
        h2 {{ color: #34495e; margin-top: 20px; font-size: 15pt; }}
        h3 {{ color: #7f8c8d; font-size: 13pt; }}
        table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
        th {{ background-color: #3498db; color: white; font-weight: bold; }}
        tr:nth-child(even) {{ background-color: #f2f2f2; }}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""
                
                css = CSS(string="@page { size: A4; margin: 2cm; }")
                html = HTML(string=html_doc, encoding='utf-8')
                html.write_pdf(str(filepath), stylesheets=[css])
                
                logger.info(f"PDF 文档已保存到 (weasyprint): {filepath}")
                return str(filepath)
                
            except ImportError:
                error_msg = "缺少 PDF 生成库。请安装以下之一：\n1. imgkit + wkhtmltopdf (推荐): pip install imgkit\n   并下载 wkhtmltopdf: https://wkhtmltopdf.org/\n2. weasyprint: pip install weasyprint (Windows 需要 GTK+)"
                logger.error(error_msg)
                raise ImportError(error_msg)
            except Exception as e:
                logger.error(f"weasyprint 导出失败: {e}", exc_info=True)
                raise
                
        except Exception as e:
            logger.error(f"导出 PDF 文档失败: {e}", exc_info=True)
            raise
